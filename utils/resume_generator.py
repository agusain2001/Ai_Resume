import os
import subprocess
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import Dict, Any, Tuple
from utils.template_manager import TemplateManager
from docx2pdf import convert

class ResumeGenerator:
    """Generate resumes in Word and PDF formats"""
    
    def __init__(self):
        self.template_manager = TemplateManager()
        self.output_dir = "output"
        os.makedirs(self.output_dir, exist_ok=True)
    
    def generate(self, resume_data: Dict[str, Any], template_name: str) -> Tuple[str, str]:
        """Generate both Word and PDF versions of the resume"""
        
        # Generate Word document
        docx_path = self._generate_docx(resume_data, template_name)
        
        # Generate PDF (from LaTeX or convert DOCX)
        pdf_path = self._generate_pdf(resume_data, template_name, docx_path)
        
        return docx_path, pdf_path
    
    def _generate_docx(self, resume_data: Dict[str, Any], template_name: str) -> str:
        """Generate Word document"""
        
        doc = Document()
        
        # Set margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.5)
            section.right_margin = Inches(0.5)
        
        # Add content based on template style
        if template_name == "template2":
            self._add_modern_style(doc, resume_data)
        else:
            self._add_professional_style(doc, resume_data)
        
        # Save document
        output_path = os.path.join(self.output_dir, "resume.docx")
        doc.save(output_path)
        
        return output_path
    
    def _add_professional_style(self, doc: Document, resume_data: Dict[str, Any]):
        """Add content in professional style"""
        
        # Header - Name and Contact Info
        personal = resume_data.get('personal_info', {})
        
        # Name
        name_para = doc.add_paragraph()
        name_run = name_para.add_run(personal.get('name', ''))
        name_run.font.size = Pt(24)
        name_run.font.bold = True
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Contact Info
        contact_info = []
        if personal.get('phone'):
            contact_info.append(personal['phone'])
        if personal.get('email'):
            contact_info.append(personal['email'])
        if personal.get('linkedin'):
            contact_info.append(personal['linkedin'])
        if personal.get('github'):
            contact_info.append(personal['github'])
        
        contact_para = doc.add_paragraph(' | '.join(contact_info))
        contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_para.runs[0].font.size = Pt(10)
        
        doc.add_paragraph()  # Spacing
        
        # Professional Summary
        if resume_data.get('summary'):
            self._add_section_heading(doc, "PROFESSIONAL SUMMARY")
            summary_para = doc.add_paragraph(resume_data['summary'])
            summary_para.runs[0].font.size = Pt(11)
            doc.add_paragraph()
        
        # Education
        if resume_data.get('education'):
            self._add_section_heading(doc, "EDUCATION")
            for edu in resume_data['education']:
                # Degree and Institution
                edu_para = doc.add_paragraph()
                degree_run = edu_para.add_run(edu.get('degree', ''))
                degree_run.font.bold = True
                degree_run.font.size = Pt(11)
                
                edu_para.add_run(' - ')
                inst_run = edu_para.add_run(edu.get('institution', ''))
                inst_run.font.size = Pt(11)
                
                # Date and GPA
                details_para = doc.add_paragraph()
                details_para.add_run(f"{edu.get('graduation_date', '')}  |  GPA: {edu.get('gpa', 'N/A')}")
                details_para.runs[0].font.size = Pt(10)
                details_para.runs[0].font.italic = True
                
            doc.add_paragraph()
        
        # Work Experience
        if resume_data.get('experience'):
            self._add_section_heading(doc, "WORK EXPERIENCE")
            for exp in resume_data['experience']:
                # Job Title and Company
                exp_para = doc.add_paragraph()
                title_run = exp_para.add_run(exp.get('title', ''))
                title_run.font.bold = True
                title_run.font.size = Pt(11)
                
                exp_para.add_run(' - ')
                company_run = exp_para.add_run(exp.get('company', ''))
                company_run.font.size = Pt(11)
                
                # Dates
                date_para = doc.add_paragraph()
                date_para.add_run(f"{exp.get('start_date', '')} - {exp.get('end_date', '')}")
                date_para.runs[0].font.size = Pt(10)
                date_para.runs[0].font.italic = True
                
                # Responsibilities
                responsibilities = exp.get('responsibilities', '')
                if responsibilities:
                    for line in responsibilities.split('\n'):
                        line = line.strip()
                        if line:
                            # Remove existing bullet points
                            line = line.replace('•', '').replace('-', '').replace('*', '').strip()
                            if line:
                                bullet_para = doc.add_paragraph(line, style='List Bullet')
                                bullet_para.runs[0].font.size = Pt(10)
                
                doc.add_paragraph()
        
        # Skills
        if resume_data.get('skills'):
            self._add_section_heading(doc, "TECHNICAL SKILLS")
            skills = resume_data['skills']
            
            if isinstance(skills, dict):
                if skills.get('technical'):
                    tech_para = doc.add_paragraph()
                    tech_para.add_run('Technical Skills: ').font.bold = True
                    tech_para.add_run(skills['technical'])
                    tech_para.runs[0].font.size = Pt(11)
                    tech_para.runs[1].font.size = Pt(11)
                
                if skills.get('soft'):
                    soft_para = doc.add_paragraph()
                    soft_para.add_run('Soft Skills: ').font.bold = True
                    soft_para.add_run(skills['soft'])
                    soft_para.runs[0].font.size = Pt(11)
                    soft_para.runs[1].font.size = Pt(11)
            else:
                skills_para = doc.add_paragraph(str(skills))
                skills_para.runs[0].font.size = Pt(11)
            
            doc.add_paragraph()
        
        # Projects
        if resume_data.get('projects'):
            self._add_section_heading(doc, "PROJECTS")
            for proj in resume_data['projects']:
                # Project Name
                proj_para = doc.add_paragraph()
                name_run = proj_para.add_run(proj.get('name', ''))
                name_run.font.bold = True
                name_run.font.size = Pt(11)
                
                # Technologies
                if proj.get('technologies'):
                    proj_para.add_run(f" | {proj['technologies']}")
                    proj_para.runs[1].font.italic = True
                    proj_para.runs[1].font.size = Pt(10)
                
                # Description
                if proj.get('description'):
                    desc_para = doc.add_paragraph(proj['description'])
                    desc_para.runs[0].font.size = Pt(10)
                
                doc.add_paragraph()
    
    def _add_modern_style(self, doc: Document, resume_data: Dict[str, Any]):
        """Add content in modern style with color accents"""
        
        # Similar to professional but with color
        personal = resume_data.get('personal_info', {})
        
        # Name with color
        name_para = doc.add_paragraph()
        name_run = name_para.add_run(personal.get('name', ''))
        name_run.font.size = Pt(26)
        name_run.font.bold = True
        name_run.font.color.rgb = RGBColor(0, 102, 204)  # Blue color
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Contact Info
        contact_info = []
        if personal.get('phone'):
            contact_info.append(personal['phone'])
        if personal.get('email'):
            contact_info.append(personal['email'])
        if personal.get('linkedin'):
            contact_info.append('LinkedIn')
        if personal.get('github'):
            contact_info.append('GitHub')
        
        contact_para = doc.add_paragraph(' | '.join(contact_info))
        contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_para.runs[0].font.size = Pt(10)
        
        doc.add_paragraph()
        
        # Use the same structure as professional but with colored headings
        self._add_professional_style(doc, resume_data)
    
    def _add_section_heading(self, doc: Document, heading_text: str):
        """Add a section heading"""
        heading = doc.add_heading(heading_text, level=2)
        heading.runs[0].font.size = Pt(14)
        heading.runs[0].font.color.rgb = RGBColor(0, 0, 0)
        heading.runs[0].font.bold = True
        
        # Add horizontal line
        doc.add_paragraph('_' * 80)
    
    def _generate_pdf(self, resume_data: Dict[str, Any], template_name: str, docx_path: str) -> str:
        """Generate PDF version"""
        
        # Method 1: Try to compile LaTeX (if available)
        try:
            return self._generate_pdf_from_latex(resume_data, template_name)
        except Exception as e:
            print(f"LaTeX compilation not available: {e}")
        
        # Method 2: Try to convert DOCX to PDF (if library available)
        try:
            return self._convert_docx_to_pdf(docx_path)
        except Exception as e:
            print(f"DOCX to PDF conversion not available: {e}")
        
        # Method 3: Create a simple text-based PDF
        return self._create_simple_pdf(resume_data)
    
    def _generate_pdf_from_latex(self, resume_data: Dict[str, Any], template_name: str) -> str:
        """Generate PDF from LaTeX template"""
        
        # Fill LaTeX template
        latex_content = self.template_manager.fill_template(template_name, resume_data)
        
        # Save LaTeX file
        latex_path = os.path.join(self.output_dir, "resume.tex")
        with open(latex_path, 'w', encoding='utf-8') as f:
            f.write(latex_content)
        
        # Compile LaTeX to PDF
        subprocess.run(
            ['pdflatex', '-output-directory', self.output_dir, latex_path],
            check=True,
            capture_output=True
        )
        
        pdf_path = os.path.join(self.output_dir, "resume.pdf")
        return pdf_path
    
    def _convert_docx_to_pdf(self, docx_path: str) -> str:
        """Convert DOCX to PDF using available tools"""
        
        pdf_path = os.path.join(self.output_dir, "resume.pdf")
        
        # Try using python-docx-to-pdf or similar
        try:
            convert(docx_path, pdf_path)
            return pdf_path
        except ImportError:
            raise Exception("PDF conversion library not available")
    
    def _create_simple_pdf(self, resume_data: Dict[str, Any]) -> str:
        """Create a simple PDF using reportlab"""
        
        try:
            from reportlab.lib.pagesizes import letter
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import inch
            
            pdf_path = os.path.join(self.output_dir, "resume.pdf")
            doc = SimpleDocTemplate(pdf_path, pagesize=letter)
            
            styles = getSampleStyleSheet()
            story = []
            
            # Add content
            personal = resume_data.get('personal_info', {})
            
            # Name
            title_style = ParagraphStyle('CustomTitle', parent=styles['Heading1'], fontSize=24, alignment=1)
            story.append(Paragraph(personal.get('name', ''), title_style))
            story.append(Spacer(1, 0.2*inch))
            
            # Contact
            contact_info = f"{personal.get('email', '')} | {personal.get('phone', '')}"
            story.append(Paragraph(contact_info, styles['Normal']))
            story.append(Spacer(1, 0.3*inch))
            
            # Summary
            if resume_data.get('summary'):
                story.append(Paragraph('<b>PROFESSIONAL SUMMARY</b>', styles['Heading2']))
                story.append(Paragraph(resume_data['summary'], styles['Normal']))
                story.append(Spacer(1, 0.2*inch))
            
            doc.build(story)
            return pdf_path
            
        except ImportError:
            # If reportlab not available, just copy the DOCX
            pdf_path = os.path.join(self.output_dir, "resume.pdf")
            # Create a placeholder file
            with open(pdf_path, 'w') as f:
                f.write("PDF generation not available. Please use the DOCX file.")
            return pdf_path